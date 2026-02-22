"""
Word document builder.

Assembles a list of per-page extracted text strings into a .docx file
with a hard page break between each original PDF page.
"""

from typing import List

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def _add_page_break(doc: Document) -> None:
    """Insert a hard Word page break (not a section break)."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def build_docx(
    page_texts: List[str],
    output_path: str,
    source_filename: str = "handwritten_document",
) -> str:
    """
    Build a .docx file from per-page extracted text strings.

    Each page's text is added as paragraphs (one per line). A hard page
    break separates consecutive pages to match the original PDF layout.

    Args:
        page_texts:       List of text strings, one per original PDF page
        output_path:      Absolute path where the .docx file should be saved
        source_filename:  Used as the document title in metadata

    Returns:
        output_path (str)

    Raises:
        ValueError: If page_texts is empty
    """
    if not page_texts:
        raise ValueError("page_texts must not be empty")

    doc = Document()
    doc.core_properties.title = f"OCR: {source_filename}"
    doc.core_properties.author = "Handwritten OCR App"

    for page_index, text in enumerate(page_texts):
        lines = text.splitlines()
        for line in lines:
            doc.add_paragraph(line)

        # Hard page break between pages (not after the last page)
        if page_index < len(page_texts) - 1:
            _add_page_break(doc)

    doc.save(output_path)
    return output_path
